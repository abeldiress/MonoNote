from summarize import generate_summary
from requests.exceptions import ConnectionError
from docx import Document
from docx.shared import Inches
import bs4, requests
from datetime import datetime
import urllib
import io
import os
import PyPDF2

try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML

import zipfile

WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA = WORD_NAMESPACE + 'p'
TEXT = WORD_NAMESPACE + 't'

headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.129 Safari/537.36'}


def extractText(k, _type):
    final = list()
    if _type == 'link':
        try:
            k_content = requests.get('http://' + k if 'htt' not in k else k, headers=headers, timeout=10).text
            d = bs4.BeautifulSoup(k_content, 'html.parser')
        except ConnectionError:
            pass
        except TimeoutError:
            pass
  
        texts = d.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
  
        # find all p in html over 26 words
        for part in texts:

            # finds all paragraphs with more than or equal to 15 words
            if part.name == 'p':
                x = str(part.text).split(' ')
                if len(x) >= 15:
                    final.append(part.text)

            # finds all title less than 9 words
            else:
                x = str(part.text).split(' ')
                if len(x) < 9:
                    final.append('\nTitle: ' + part.text + '\n')

        # repeats the below comment until it goes through list without removing any title
        tick = True    
        while tick:
            tick = False

            # removes all titles that have no content
            for i in range(len(final)):
                try:
                    if 'Title: ' in final[i] and len(str(final[i-1]).split(' ')) < 25: 
                        tick = True           
                        del final[i-1]
                except IndexError:
                    pass
    elif _type == 'text':
        final = k.split('\n')

        for partIndex in range(len(final) - 1):
            # finds all paragraphs more than 15 words
            try:
                if final[partIndex] == '' or len(final[partIndex].split(' ')) < 15:
                    del final[partIndex]
            except:
                break
            
    elif _type == 'docx':
        document = zipfile.ZipFile(k)
        xml_content = document.read('word/document.xml')
        document.close()
        tree = XML(xml_content)

        paragraphs = []
        for paragraph in tree.getiterator(PARA):
            texts = [node.text
                    for node in paragraph.getiterator(TEXT)
                    if node.text]
            if texts:
                paragraphs.append(''.join(texts))

        final = [para for para in paragraphs if not len(para.split(' ')) < 10]

    elif _type == 'pdf':
        pdfFileObj = open(k, 'rb')
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

        k = str()
        for i in range(pdfReader.numPages):
            k = k + pdfReader.getPage(i).extractText()
            print(pdfReader.getPage(i).extractText())
        
        final = k.split('\n')

        for partIndex in range(len(final)):
            # finds all paragraphs more than 15 words
            if final[partIndex] == '' or len(final[partIndex].split(' ')) < 15:
                if 'Title: ' not in final[partIndex]:
                    delTick = True
                    del final[partIndex]
                    break
        
        pdfFileObj.close()
                
    # this removes any unecessary indents, tabs, or enters
    for partIndex in range(len(final)):
        final[partIndex] = final[partIndex].replace('\n', '')
        final[partIndex] = final[partIndex].replace('\r', '')
        final[partIndex] = final[partIndex].replace('\t', '')

    return final

def summarize(final, percentage):
    n = final
    sum_num = float(percentage)

    for paragraph in range(len(n)):
        if 'Title: ' not in n[paragraph]:
            sentence_length = len(n[paragraph].split('. '))
            print(sentence_length * (sum_num/100))
            new_length = round(sentence_length * (sum_num/100))
            
            summarized = generate_summary(n[paragraph], new_length)
            n[paragraph] = summarized
            
            print('Sentence Length: ' + str(sentence_length))
            print('New Length: ' + str(new_length))
                    
            print('--------------------------------------------')

    print('Text Summarized!')

    return n

def get_images(summary, topic):
    for partIndex in range(len(summary) - 1):
        if 'Title: ' in summary[partIndex]:
            try:
                summary_images = bs4.BeautifulSoup(requests.get('https://www.google.com/search?q=' + summary[partIndex] + ' ' + topic + '&tbm=isch&hl=en&hl=en&tbs=isz%3Al').content, 'html.parser')         
                image = summary_images.select('img[style="border:1px solid #ccc;padding:1px"]')[0]
                summary.insert(partIndex + 2, 'Image%%:' + str(image['src']).replace('http', 'https'))
            except:
                continue
    
    return summary

def get_links(summary, topic):
    summary.append('\nTitle: Related Links\n')
    related_page = bs4.BeautifulSoup(requests.get('https://www.google.com/search?q=' + topic).text, 'html.parser')
    related_links = related_page.select('.kCrYT > a')
    
    for link in related_links[5:len(related_links) - 1]:
        summary.append(str(link['href']).replace('/url?q=', '').split('&sa')[0] + '\n')
        #print(link['href'])
    
    return summary

def download_as_doc(n, topic, date_created, id):
    sum_doc = Document()
    sum_doc.add_heading(topic, level=0)
    sum_doc.add_heading('Date created: ' + date_created, level=1)
    for part in n:
        if 'Title: ' in part:
            t = sum_doc.add_heading(part.replace('Title: ', ''), level=2)
        
        elif 'Image%%:' in part:
            i = sum_doc.add_picture(io.BytesIO(urllib.request.urlopen(part.replace('Image%%:', '')).read()))
        else:
            sum_doc.add_paragraph(part)

    # makes a file, reads file bytes, then deletes file
    sum_doc.save('{}.docx'.format(id))
    t = open('{}.docx'.format(id), 'rb').read()
    os.remove('{}.docx'.format(id))

    return t

def download_as_text(n, topic, date_created, id):
    full_note = str()
    full_note = full_note + topic + '\n\n\n'
    full_note = full_note + 'Date created: ' + date_created + '\n\n'

    for part in n:
        if 'Title: ' in part:
            full_note = full_note + part.replace('Title: ', '') + '\n'
        else:
            full_note = full_note + part + '\n\n'
    
    t = bytes(full_note, encoding='utf-8')
    return t

def download_as_html(n, topic, date_created, id):
    full_note = '<title>{}</title><link href="https://fonts.googleapis.com/css?family=Catamaran:100,200,300,400,500,600,700,800,900" rel="stylesheet"><link href="https://fonts.googleapis.com/css?family=Lato:100,100i,300,300i,400,400i,700,700i,900,900i" rel="stylesheet"><div style="margin: 0; width: 80%; padding-left: 4%; padding-top: 2%;">'.format(topic)
    full_note = full_note + '<h1 style="font-family: Catamaran, sans-serif; font-size: 2.5rem; margin: 0;">{}</h1>'.format(topic)
    full_note = full_note + '<h4 style="font-family: Lato, sans-serif; color: #888; margin: 0; padding-top: 1%;">Made on {}</h4>'.format(date_created)

    for part in n:
        if 'Title: ' in part:
            full_note = full_note + '<h3 style="font-family: Catamaran, sans-serif; font-size: 1.3rem;">{}</h3>'.format(part.replace('Title: ', ''))
        elif 'Image%%:' in part:
            full_note = full_note + '<img src="{}">'.format(part.replace('Image%%:', ''))
        else:
            full_note = full_note + '<p style="font-family: Lato, sans-serif;">{}</p>'.format(part)
    
    full_note = full_note + '</div>'

    t = bytes(full_note, encoding='utf-8')

    return t

def download_as_pdf(n, topic, date_created, id):
    full_note = '<title>{}</title><link href="https://fonts.googleapis.com/css?family=Catamaran:100,200,300,400,500,600,700,800,900" rel="stylesheet"><link href="https://fonts.googleapis.com/css?family=Lato:100,100i,300,300i,400,400i,700,700i,900,900i" rel="stylesheet"><div style="margin: 0; width: 80%; padding-left: 4%; padding-top: 2%;">'.format(topic)
    full_note = full_note + '<h1 style="font-family: Catamaran, sans-serif; font-size: 2.5rem; margin: 0;">{}</h1>'.format(topic)
    full_note = full_note + '<h4 style="font-family: Lato, sans-serif; color: #888; margin: 0; padding-top: 1%;">Made on {}</h4>'.format(date_created)

    for part in n:
        if 'Title: ' in part:
            full_note = full_note + '<h3 style="font-family: Catamaran, sans-serif; font-size: 1.3rem;">{}</h3>'.format(part.replace('Title: ', ''))
        elif 'Image%%:' in part:
            full_note = full_note + '<img src="{}">'.format(part.replace('Image%%:', ''))
        else:
            full_note = full_note + '<p style="font-family: Lato, sans-serif;">{}</p>'.format(part)
    
    full_note = full_note + '</div>'

    t = requests.post(
        'https://api.html2pdf.app/v1/generate?apiKey=67c7898e82e38587e0e1805a152312138970bcd255a996034da9918a7ccb923b',
        data={'html': full_note}
    ).content

    return t